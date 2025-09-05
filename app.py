# # """
# # Literature Snapshot API — Multi-source, OA-first, typo-tolerant search
# # List format with per-paper summaries and inline clickable citations
# # Exact PDF page count control (supports large values, e.g., 50+ pages)
# # Temp folder lifecycle for all downloaded files (auto-cleanup when not archived)
# #
# # Key additions in this version
# # - All transient downloads (e.g., source PDFs used for snippet/image extraction) go into a per-request temp folder
# #   under TEMP_ROOT_DIR (defaults to OS temp, e.g., /tmp/literature_snapshot or %TEMP%\literature_snapshot).
# # - If archive=0 (default behavior you want), the per-request temp folder is deleted at the end of the request,
# #   so nothing remains on disk once the PDF/DOCX has been generated and streamed to the user.
# # - If archive=1, files are saved under ./local/{query_slug}/ and NOT deleted (plus manifest.json is written).
# #
# # Endpoint
# # - GET /report?query=...&max_pages=50&file=pdf
# #   Parameters:
# #     - query       (required): topic (typos tolerated)
# #     - max_pages   (required for exact pages): integer 1–100 (supports 50+)
# #     - file        (optional): pdf (default) or docx
# #     - max_results (optional): 10–500 (default 120)
# #     - sources     (optional): comma-separated in [openalex,arxiv,eupmc,crossref,s2] (default all)
# #     - since       (optional): min year (default 2015)
# #     - archive     (optional): 1 or 0 (default 0 here) — when 0, temp files are cleaned up after response
# #
# # Notes
# # - For exact page counts, ensure pypdf is installed (see requirements.txt).
# # - Inline citations are clickable links that open the source page/PDF in a new tab.
# # """
# #
# # from __future__ import annotations
# #
# import io
# import os
# import re
# import math
# import html
# import json
# import time
# import uuid
# import shutil
# import logging
# import tempfile
# from datetime import datetime
# from typing import List, Dict, Any, Optional, Tuple
#
# import requests
# from flask import Flask, request, send_file, jsonify
# from flask_cors import CORS  # CORS for browser fetches
#
# # PDF generation
# from reportlab.lib.pagesizes import A4
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib import utils
#
# # PDF page counting (to guarantee exact pages)
# try:
#     from pypdf import PdfReader
#     PYPDF_AVAILABLE = True
# except Exception:
#     PYPDF_AVAILABLE = False
#
# # Optional DOCX output (exact pages not guaranteed in DOCX)
# try:
#     from docx import Document as DocxDocument
#     from docx.shared import Inches, Pt
#     from docx.oxml import OxmlElement
#     from docx.oxml.ns import qn
#     DOCX_AVAILABLE = True
# except Exception:
#     DOCX_AVAILABLE = False
#
# # Optional PyMuPDF for image/snippet extraction
# try:
#     import fitz  # type: ignore
#     PYMUPDF_AVAILABLE = True
# except Exception:
#     PYMUPDF_AVAILABLE = False
#
# # Optional scikit-learn for TF-IDF ranking
# try:
#     from sklearn.feature_extraction.text import TfidfVectorizer
#     from sklearn.metrics.pairwise import cosine_similarity
#     SKLEARN_AVAILABLE = True
# except Exception:
#     SKLEARN_AVAILABLE = False
#
# # ------------------------------------------------------------------------------
# # Config and constants
# # ------------------------------------------------------------------------------
#
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)
#
# APP_TITLE = "Literature Snapshot"
#
# DEFAULT_MAX_PAGES = 2
# DEFAULT_MAX_RESULTS = 120   # higher default to support large page counts
# DEFAULT_SINCE = 2015
# AUTHORIZED_SOURCES = {"openalex", "arxiv", "eupmc", "crossref", "s2"}
#
# # External APIs
# OPENALEX_BASE = "https://api.openalex.org/works"
# ARXIV_BASE = "http://export.arxiv.org/api/query"
# EUROPE_PMC_BASE = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
# CROSSREF_BASE = "https://api.crossref.org/works"
# S2_BASE = "https://api.semanticscholar.org/graph/v1/paper/search"
#
# # Character budget heuristic per page (tuned for list format and 10pt body)
# CHAR_BUDGET_PER_PAGE = 2200
#
# # Representative image filter
# MIN_IMG_WIDTH = 300
# MIN_IMG_HEIGHT = 150
# MIN_IMG_SIZE_BYTES = 20 * 1024
#
# # Archive root for persisted PDFs and manifests (when archive=1)
# LOCAL_ROOT_DIR = "./local"
#
#
#
#
#
# # Temp root for per-request transient files
# TEMP_ROOT_DIR = os.environ.get("TEMP_ROOT_DIR") or os.path.join(tempfile.gettempdir(), "literature_snapshot")
#
#
# # Default archiving toggle for this build: do NOT keep files unless explicitly requested
# ARCHIVE_PDFS_DEFAULT = False
#
# # ------------------------------------------------------------------------------
# # App bootstrap (CORS)
# # ------------------------------------------------------------------------------
#
# app = Flask(__name__)
# CORS(
#     app,
#     resources={r"/*": {"origins": "*"}},
#     expose_headers=["Content-Disposition"]
# )
# os.makedirs(LOCAL_ROOT_DIR, exist_ok=True)
# os.makedirs(TEMP_ROOT_DIR, exist_ok=True)
#
# def detect_multiprocessing_availability() -> bool:
#     try:
#         import _multiprocessing  # noqa: F401
#         return True
#     except Exception:
#         return False
#
# def flask_run_options(desired_debug: bool) -> Dict[str, Any]:
#     if desired_debug and detect_multiprocessing_availability():
#         return {"debug": True, "use_reloader": True}
#     else:
#         return {"debug": False, "use_reloader": False}
#
# # ------------------------------------------------------------------------------
# # Utilities
# # ------------------------------------------------------------------------------
#
# def normspace(s: str) -> str:
#     return re.sub(r"\s+", " ", s or "").strip()
#
# def strip_html(jats_or_html: str) -> str:
#     no_tags = re.sub(r"<[^>]+>", " ", jats_or_html or "")
#     return normspace(html.unescape(no_tags))
#
# def sentence_case(title: str) -> str:
#     if not title:
#         return title
#     t = title.strip()
#     return t[:1].upper() + t[1:]
#
# def year_from_date(d: Optional[str]) -> Optional[int]:
#     if not d:
#         return None
#     m = re.match(r"(\d{4})", d)
#     return int(m.group(1)) if m else None
#
# def html_escape(s: str) -> str:
#     return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
#
# def safe_request_json(url: str, params: Dict[str, Any], timeout: int = 30, headers: Optional[Dict[str, str]] = None) -> Optional[Dict[str, Any]]:
#     try:
#         r = requests.get(url, params=params, timeout=timeout, headers=headers or {})
#         r.raise_for_status()
#         return r.json()
#     except Exception as e:
#         logger.warning("Request failed: %s params=%s error=%s", url, params, e)
#         return None
#
# def safe_slug(s: str) -> str:
#     s = normspace(s).lower()
#     s = re.sub(r"[^\w\-]+", "_", s)
#     s = re.sub(r"_+", "_", s).strip("_")
#     return s[:80] or f"query_{int(time.time())}"
#
# def safe_filename(s: str) -> str:
#     s = normspace(s)
#     s = re.sub(r"[^\w\-]+", "_", s)
#     return s[:120] or f"paper_{int(time.time())}"
#
# # ------------------------------------------------------------------------------
# # Temp and archive dirs
# # ------------------------------------------------------------------------------
#
# def make_request_temp_dir() -> str:
#     """
#     Create a unique temp folder for this request under TEMP_ROOT_DIR.
#     """
#     uid = uuid.uuid4().hex[:12]
#     path = os.path.join(TEMP_ROOT_DIR, f"req_{uid}")
#     os.makedirs(path, exist_ok=True)
#     return path
#
# def cleanup_dir(path: Optional[str]) -> None:
#     if not path:
#         return
#     try:
#         shutil.rmtree(path, ignore_errors=True)
#         logger.info("Cleaned temp folder: %s", path)
#     except Exception as e:
#         logger.warning("Failed to clean temp folder %s: %s", path, e)
#
# def get_archive_dir(query: str) -> str:
#     slug = safe_slug(query)
#     path = os.path.join(LOCAL_ROOT_DIR, slug)
#     os.makedirs(path, exist_ok=True)
#     return path
#
# # ------------------------------------------------------------------------------
# # Query normalization for typo tolerance
# # ------------------------------------------------------------------------------
#
# def query_variants(q: str) -> List[str]:
#     """
#     Return a list of query variants to tolerate simple spelling/formatting mistakes.
#     """
#     q = q or ""
#     base = normspace(q)
#     lower = base.lower()
#     collapsed = re.sub(r"(.)\1{2,}", r"\1\1", lower)
#     punct_trim = re.sub(r"[\"'`~^|]+", " ", lower)
#     alnum = re.sub(r"[^\w\s]+", " ", lower)
#     seen, out = set(), []
#     for v in [base, lower, collapsed, punct_trim, alnum]:
#         v = normspace(v)
#         if v and v not in seen:
#             seen.add(v); out.append(v)
#     return out
#
# # ------------------------------------------------------------------------------
# # Record helpers and inline citations
# # ------------------------------------------------------------------------------
#
# def unify_record(
#     source: str,
#     title: Optional[str],
#     abstract: Optional[str],
#     year: Optional[int],
#     url: Optional[str],
#     pdf_url: Optional[str],
#     doi: Optional[str],
#     authors: Optional[List[str]] = None,
#     venue: Optional[str] = None,
# ) -> Dict[str, Any]:
#     return {
#         "source": source,
#         "title": normspace(title or ""),
#         "abstract": normspace(abstract or ""),
#         "year": year,
#         "url": url,
#         "pdf_url": pdf_url,
#         "doi": (doi or "").lower() or None,
#         "authors": authors or [],
#         "venue": normspace(venue or ""),
#         "_local_pdf": None,
#         "_local_pdf_rel": None,
#         "_snippet": None,
#     }
#
# def is_valid_record(rec: Dict[str, Any]) -> bool:
#     if not rec["title"] or not rec["abstract"]:
#         return False
#     if rec["year"] and rec["year"] < DEFAULT_SINCE:
#         return False
#     return True
#
# def dedupe_records(recs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
#     seen = set()
#     out = []
#     for r in recs:
#         key = r.get("doi") or r["title"].lower()
#         if key in seen:
#             continue
#         seen.add(key)
#         out.append(r)
#     return out
#
# def last_names(authors: List[str]) -> List[str]:
#     ln = []
#     for a in authors or []:
#         parts = [p for p in a.replace(",", " ").split() if p]
#         if parts:
#             ln.append(parts[-1])
#     return ln
#
# def inline_citation_label(rec: Dict[str, Any]) -> str:
#     ln = last_names(rec.get("authors") or [])
#     y = rec.get("year")
#     ystr = str(y) if y else "n.d."
#     if not ln:
#         return f"({ystr})"
#     if len(ln) == 1:
#         return f"({ln[0]}, {ystr})"
#     if len(ln) == 2:
#         return f"({ln[0]} & {ln[1]}, {ystr})"
#     return f"({ln[0]} et al., {ystr})"
#
# def best_href(rec: Dict[str, Any]) -> Optional[str]:
#     return rec.get("pdf_url") or rec.get("url")
#
# # ------------------------------------------------------------------------------
# # Providers
# # ------------------------------------------------------------------------------
#
# def reconstruct_openalex_abstract(inv_index: Optional[Dict[str, List[int]]]) -> Optional[str]:
#     if not inv_index:
#         return None
#     max_pos = -1
#     for positions in inv_index.values():
#         if positions:
#             max_pos = max(max_pos, max(positions))
#     if max_pos < 0:
#         return None
#     tokens: List[Optional[str]] = [None] * (max_pos + 1)
#     for token, positions in inv_index.items():
#         for pos in positions:
#             if 0 <= pos < len(tokens):
#                 tokens[pos] = token
#     return normspace(" ".join(t or "" for t in tokens))
#
# def fetch_openalex(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
#     params = {
#         "filter": f"concepts.display_name.search:{query},display_name.search:{query},has_fulltext:true,has_abstract:true,from_publication_date:{since}-01-01",
#         "sort": "relevance_score:desc,cited_by_count:desc",
#         "per-page": min(limit, 25),
#         "page": 1,
#     }
#     data = safe_request_json(OPENALEX_BASE, params)
#     results = data.get("results", []) if data else []
#     out: List[Dict[str, Any]] = []
#     for w in results:
#         abstract = w.get("abstract")
#         if not abstract and w.get("abstract_inverted_index"):
#             abstract = reconstruct_openalex_abstract(w["abstract_inverted_index"])
#         title = w.get("title")
#         year = w.get("publication_year")
#         doi = w.get("doi")
#         venue = (w.get("host_venue") or {}).get("display_name")
#         authors = []
#         for a in (w.get("authorships") or []):
#             nm = (a.get("author") or {}).get("display_name")
#             if nm: authors.append(nm)
#         pdf_url = None
#         try:
#             best = w.get("best_oa_location") or {}
#             if best.get("pdf_url"): pdf_url = best["pdf_url"]
#         except Exception:
#             pass
#         if not pdf_url:
#             for loc in (w.get("oa_locations") or []):
#                 if loc.get("pdf_url"): pdf_url = loc["pdf_url"]; break
#         landing = (w.get("primary_location") or {}).get("landing_page_url")
#         url = landing or (w.get("host_venue") or {}).get("url") or w.get("id")
#         rec = unify_record("OPENALEX", title, abstract, year, url, pdf_url, doi, authors, venue)
#         if is_valid_record(rec): out.append(rec)
#         if len(out) >= limit: break
#     return out
#
# def fetch_arxiv(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
#     per_page = min(limit, 25)
#     params = {
#         "search_query": f'all:{query}',
#         "start": 0,
#         "max_results": per_page,
#         "sortBy": "relevance",
#         "sortOrder": "descending",
#     }
#     try:
#         r = requests.get(ARXIV_BASE, params=params, timeout=30)
#         r.raise_for_status()
#         import xml.etree.ElementTree as ET
#         root = ET.fromstring(r.text)
#         ns = {"a": "http://www.w3.org/2005/Atom"}
#         out: List[Dict[str, Any]] = []
#         for entry in root.findall("a:entry", ns):
#             title = normspace(entry.findtext("a:title", default="", namespaces=ns))
#             abstract = normspace(entry.findtext("a:summary", default="", namespaces=ns))
#             published = entry.findtext("a:published", default="", namespaces=ns)
#             year = year_from_date(published)
#             if year and year < since: continue
#             url = entry.findtext("a:id", default="", namespaces=ns)
#             pdf_url = None
#             for link in entry.findall("a:link", ns):
#                 if link.attrib.get("type") == "application/pdf":
#                     pdf_url = link.attrib.get("href"); break
#             authors = [normspace(a.findtext("a:name", default="", namespaces=ns)) for a in entry.findall("a:author", ns)]
#             venue = "arXiv"
#             rec = unify_record("ARXIV", title, abstract, year, url, pdf_url, None, authors, venue)
#             if is_valid_record(rec): out.append(rec)
#             if len(out) >= limit: break
#         return out
#     except Exception as e:
#         logger.warning("arXiv fetch failed: %s", e)
#         return []
#
# def fetch_europe_pmc(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
#     params = {
#         "query": f'({query}) AND PUB_YEAR:[{since} TO 3000] AND HAS_ABSTRACT:Y',
#         "resultType": "core",
#         "pageSize": min(limit, 25),
#         "format": "json",
#     }
#     data = safe_request_json(EUROPE_PMC_BASE, params)
#     results = data.get("resultList", {}).get("result", []) if data else []
#     out: List[Dict[str, Any]] = []
#     for it in results:
#         title = it.get("title")
#         abstract = it.get("abstractText")
#         year = int(it["pubYear"]) if it.get("pubYear") and str(it["pubYear"]).isdigit() else None
#         if year and year < since: continue
#         doi = it.get("doi")
#         url = None
#         if it.get("fullTextUrlList") and it["fullTextUrlList"].get("fullTextUrl"):
#             url = it["fullTextUrlList"]["fullTextUrl"][0].get("url")
#         url = url or it.get("doi")
#         pdf_url = None
#         for ft in (it.get("fullTextUrlList", {}) or {}).get("fullTextUrl", []) or []:
#             u = ft.get("url") or ""
#             if ft.get("documentStyle") == "pdf" or ("pdf" in u.lower() and ft.get("availability", "").lower().startswith("open")):
#                 pdf_url = u; break
#         venue = it.get("journalTitle") or "Europe PMC"
#         authors: List[str] = []
#         if it.get("authorList") and it["authorList"].get("author"):
#             for a in it["authorList"]["author"]:
#                 nm = a.get("fullName") or a.get("lastName") or ""
#                 if a.get("firstName"): nm = f"{a['firstName']} {nm}".strip()
#                 if nm: authors.append(nm)
#         elif it.get("authorString"):
#             authors = [x.strip() for x in re.split(r"[;,]", it["authorString"]) if x.strip()]
#         rec = unify_record("EUROPE PMC", title, abstract, year, url, pdf_url, doi, authors, venue)
#         if is_valid_record(rec): out.append(rec)
#         if len(out) >= limit: break
#     return out
#
# def fetch_crossref(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
#     params = {
#         "query.title": query,
#         "filter": f"from-pub-date:{since}-01-01,type:journal-article,has-abstract:true",
#         "rows": min(limit, 25),
#         "sort": "relevance",
#         "order": "desc",
#     }
#     try:
#         r = requests.get(CROSSREF_BASE, params=params, timeout=30)
#         r.raise_for_status()
#         data = r.json()
#         items = data.get("message", {}).get("items", []) or []
#         out: List[Dict[str, Any]] = []
#         for it in items:
#             title = normspace(" ".join(it.get("title") or []))
#             abstract_raw = it.get("abstract")
#             abstract = strip_html(abstract_raw) if abstract_raw else ""
#             year = None
#             if it.get("issued", {}).get("date-parts"):
#                 year = it["issued"]["date-parts"][0][0]
#             if year and year < since: continue
#             doi = it.get("DOI")
#             url = it.get("URL")
#             pdf_url = None
#             for link in it.get("link", []) or []:
#                 if link.get("content-type") == "application/pdf":
#                     pdf_url = link.get("URL"); break
#             venue = ""
#             if it.get("container-title"):
#                 venue = " ".join(it.get("container-title"))
#             authors: List[str] = []
#             for a in it.get("author", []) or []:
#                 given = a.get("given", ""); family = a.get("family", "")
#                 nm = " ".join([given, family]).strip() or (a.get("name") or "")
#                 if nm: authors.append(nm)
#             oa_license = any(("open" in (lic.get("URL", "") + lic.get("content-version", "")).lower()) for lic in it.get("license", []) or [])
#             rec = unify_record("CROSSREF", title, abstract, year, url, pdf_url, doi, authors, venue)
#             if is_valid_record(rec) and (pdf_url or oa_license): out.append(rec)
#             if len(out) >= limit: break
#         return out
#     except Exception as e:
#         logger.warning("Crossref fetch failed: %s", e)
#         return []
#
# def fetch_semantic_scholar(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
#     params = {
#         "query": query,
#         "fields": "title,abstract,year,url,openAccessPdf,externalIds,authors,venue",
#         "limit": min(limit, 25),
#         "offset": 0,
#     }
#     headers = {}
#     if os.environ.get("S2_API_KEY"):
#         headers["x-api-key"] = os.environ["S2_API_KEY"]
#     data = safe_request_json(S2_BASE, params, headers=headers)
#     if not data:
#         return []
#     out: List[Dict[str, Any]] = []
#     for it in data.get("data", []) or []:
#         year = it.get("year")
#         if year and year < since: continue
#         title = it.get("title") or ""
#         abstract = it.get("abstract") or ""
#         doi = (it.get("externalIds") or {}).get("DOI")
#         url = it.get("url")
#         pdf_url = (it.get("openAccessPdf") or {}).get("url")
#         authors = [a.get("name") for a in (it.get("authors") or []) if a.get("name")]
#         venue = it.get("venue") or ""
#         rec = unify_record("SEMANTIC SCHOLAR", title, abstract, year, url, pdf_url, doi, authors, venue)
#         if is_valid_record(rec): out.append(rec)
#         if len(out) >= limit: break
#     return out
#
# FETCHERS = {
#     "openalex": fetch_openalex,
#     "arxiv": fetch_arxiv,
#     "eupmc": fetch_europe_pmc,
#     "crossref": fetch_crossref,
#     "s2": fetch_semantic_scholar,
# }
#
# # ------------------------------------------------------------------------------
# # Per-request storage handling (temp vs archive)
# # ------------------------------------------------------------------------------
#
# def download_pdf_to_dir(url: str, rec: Dict[str, Any], dest_dir: str) -> Optional[str]:
#     if not url:
#         return None
#     try:
#         r = requests.get(url, timeout=60, stream=True)
#         r.raise_for_status()
#         ts = datetime.utcnow().strftime("%Y%m%d%H%M%S")
#         tag = rec.get("doi") or rec.get("title") or "paper"
#         filename = f"{rec.get('source','src')}_{safe_filename(tag)}_{ts}.pdf"
#         local_path = os.path.join(dest_dir, filename)
#         with open(local_path, "wb") as f:
#             for chunk in r.iter_content(chunk_size=8192):
#                 if chunk: f.write(chunk)
#         abs_path = os.path.abspath(local_path)
#         rec["_local_pdf"] = abs_path
#         rec["_local_pdf_rel"] = None
#         return abs_path
#     except Exception as e:
#         logger.warning("Failed to download PDF %s: %s", url, e)
#         return None
#
# def ensure_local_pdf(rec: Dict[str, Any], dest_dir: str) -> Optional[str]:
#     if rec.get("_local_pdf") and os.path.exists(rec["_local_pdf"]):
#         return rec["_local_pdf"]
#     if not rec.get("pdf_url"):
#         return None
#     return download_pdf_to_dir(rec["pdf_url"], rec, dest_dir)
#
# def write_manifest(dest_dir: str, query: str, records: List[Dict[str, Any]]) -> str:
#     payload = {
#         "query": query,
#         "generated_at_utc": datetime.utcnow().isoformat() + "Z",
#         "folder": os.path.abspath(dest_dir),
#         "records": []
#     }
#     for r in records:
#         payload["records"].append({
#             "source": r.get("source"),
#             "title": r.get("title"),
#             "authors": r.get("authors"),
#             "year": r.get("year"),
#             "venue": r.get("venue"),
#             "doi": r.get("doi"),
#             "url": r.get("url"),
#             "pdf_url": r.get("pdf_url"),
#             "local_pdf": r.get("_local_pdf"),  # absolute path in archive dir
#         })
#     manifest_path = os.path.join(dest_dir, "manifest.json")
#     with open(manifest_path, "w", encoding="utf-8") as f:
#         json.dump(payload, f, ensure_ascii=False, indent=2)
#     logger.info("Wrote manifest for query: %s", os.path.basename(dest_dir))
#     return manifest_path
#
# # ------------------------------------------------------------------------------
# # Ranking, snippets, and list-style summaries
# # ------------------------------------------------------------------------------
#
# def compute_similarity_scores(query: str, texts: List[str]) -> List[float]:
#     if SKLEARN_AVAILABLE:
#         vec = TfidfVectorizer(stop_words="english", max_features=20000)
#         X = vec.fit_transform([query] + texts)
#         sims = cosine_similarity(X[0:1], X[1:]).flatten().tolist()
#         return sims
#     # Fallback: token overlap
#     q_tokens = set(re.findall(r"\w+", query.lower()))
#     sims: List[float] = []
#     for t in texts:
#         t_tokens = set(re.findall(r"\w+", (t or "").lower()))
#         inter = len(q_tokens & t_tokens)
#         denom = math.sqrt(max(len(q_tokens), 1) * max(len(t_tokens), 1))
#         sims.append(inter / denom if denom else 0.0)
#     return sims
#
# def extract_text_snippet_from_pdf(local_pdf_path: str, max_chars: int = 1400) -> str:
#     if not PYMUPDF_AVAILABLE or not local_pdf_path:
#         return ""
#     doc = None
#     try:
#         doc = fitz.open(local_pdf_path)
#         txts: List[str] = []
#         for i in range(min(len(doc), 3)):
#             page = doc.load_page(i)
#             txts.append(page.get_text("text"))
#         snippet = normspace(" ".join(txts))[:max_chars]
#         return snippet
#     except Exception:
#         return ""
#     finally:
#         if doc:
#             doc.close()
#
# def rerank_records(query: str, records: List[Dict[str, Any]], dest_dir: str, use_pdf_snippets: bool = True, snippet_limit: int = 8) -> List[Dict[str, Any]]:
#     enriched: List[Dict[str, Any]] = []
#     for r in records[:snippet_limit]:
#         if use_pdf_snippets and r.get("pdf_url"):
#             local_pdf = ensure_local_pdf(r, dest_dir)
#             if local_pdf:
#                 snip = extract_text_snippet_from_pdf(local_pdf, max_chars=1400)
#                 if snip:
#                     r = dict(r)
#                     r["_snippet"] = snip
#         enriched.append(r)
#     enriched.extend(records[snippet_limit:])
#     texts = [f"{rr['title']} {rr['abstract']} {' ' + (rr['_snippet'] or '') if rr.get('_snippet') else ''}" for rr in enriched]
#     sims = compute_similarity_scores(query, texts)
#     ranked = [rec for _, rec in sorted(zip(sims, enriched), key=lambda x: x[0], reverse=True)]
#     return ranked
#
# def split_sentences(text: str) -> List[str]:
#     s = normspace(text)
#     if not s:
#         return []
#     parts = re.split(r'(?<=[.!?])\s+', s)
#     return [p.strip() for p in parts if p.strip()]
#
# def top_sentences_for_record(query: str, rec: Dict[str, Any], max_sents: int = 3) -> List[Dict[str, str]]:
#     base = rec.get("abstract") or ""
#     if rec.get("_snippet"):
#         base = base + " " + rec["_snippet"]
#     sents = split_sentences(base)
#     if not sents:
#         return []
#     # Score by token overlap to pick the most relevant sentences
#     q_tokens = set(re.findall(r"\w+", query.lower()))
#     scored = []
#     for s in sents:
#         tokens = set(re.findall(r"\w+", s.lower()))
#         inter = len(q_tokens & tokens)
#         denom = math.sqrt(max(len(q_tokens), 1) * max(len(tokens), 1))
#         score = inter / denom if denom else 0.0
#         scored.append((score, s))
#     scored.sort(key=lambda x: x[0], reverse=True)
#     chosen = [s for _, s in scored[:max_sents]]
#     label = inline_citation_label(rec)
#     href = best_href(rec)
#     out = []
#     for s in chosen:
#         s_clean = s.rstrip()
#         if not s_clean.endswith(('.', '!', '?')):
#             s_clean += '.'
#         out.append({"sentence": s_clean, "label": label, "href": href or ""})
#     return out
#
# def format_meta_line(rec: Dict[str, Any]) -> str:
#     parts = []
#     if rec.get("source"):
#         parts.append(rec["source"])
#     if rec.get("year"):
#         parts.append(str(rec["year"]))
#     if rec.get("venue"):
#         parts.append(rec["venue"])
#     return " • ".join(parts)
#
# # ------------------------------------------------------------------------------
# # Image extraction
# # ------------------------------------------------------------------------------
#
# def _is_good_image(pix: "fitz.Pixmap", img_bytes: bytes) -> bool:
#     return (pix.width >= MIN_IMG_WIDTH and pix.height >= MIN_IMG_HEIGHT and len(img_bytes) >= MIN_IMG_SIZE_BYTES)
#
# def extract_one_image_from_records(records: List[Dict[str, Any]], dest_dir: str) -> Optional[bytes]:
#     if not PYMUPDF_AVAILABLE:
#         return None
#     for r in records:
#         local_pdf = ensure_local_pdf(r, dest_dir)
#         if not local_pdf:
#             continue
#         doc = None
#         try:
#             doc = fitz.open(local_pdf)
#             for p in range(min(len(doc), 5)):
#                 page = doc.load_page(p)
#                 for img_info in page.get_images(full=True):
#                     xref = img_info[0]
#                     base_image = doc.extract_image(xref)
#                     img_bytes = base_image.get("image")
#                     if not img_bytes:
#                         continue
#                     pix = fitz.Pixmap(doc, xref)
#                     if pix.n >= 5:
#                         pix = fitz.Pixmap(fitz.csRGB, pix)
#                     if _is_good_image(pix, img_bytes):
#                         return img_bytes
#         except Exception as e:
#             logger.warning("Image extraction failed for %s: %s", r.get("title"), e)
#         finally:
#             if doc:
#                 doc.close()
#     return None
#
# # ------------------------------------------------------------------------------
# # Output generation — LIST FORMAT, exact pages for PDF
# # ------------------------------------------------------------------------------
#
# def _img_flowable_from_bytes(img_bytes: bytes, max_width: float) -> Image:
#     bio = io.BytesIO(img_bytes)
#     ir = utils.ImageReader(bio)
#     iw, ih = ir.getSize()
#     aspect = ih / iw
#     width = min(max_width, iw)
#     height = width * aspect
#     bio.seek(0)
#     return Image(bio, width=width, height=height)
#
# def dynamic_max_sents(target_pages: int) -> int:
#     base = 3
#     extra = min(9, max(0, (target_pages - 5) // 5 * 2))
#     return base + extra
#
# def make_flow_for_budget(query: str, records: List[Dict[str, Any]], doc_width: float, char_budget: int, target_pages: int) -> List[Any]:
#     styles = getSampleStyleSheet()
#     if "TitleSmall" not in styles:
#         styles.add(ParagraphStyle(name="TitleSmall", parent=styles['h2'], fontSize=16, leading=20, spaceAfter=12))
#         styles.add(ParagraphStyle(name="PaperTitle", parent=styles['h4'], fontSize=12, leading=15, spaceAfter=2))
#         styles.add(ParagraphStyle(name="Meta", parent=styles['Normal'], fontSize=9.5, leading=12, textColor="#555555", spaceAfter=4))
#         styles.add(ParagraphStyle(name="Body", parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=10))
#     flow: List[Any] = []
#     flow.append(Paragraph(f"{APP_TITLE}: {html_escape(query)}", styles["TitleSmall"]))
#
#     used = 0
#     max_sents_per_rec = dynamic_max_sents(target_pages)
#
#     for r in records:
#         title = r["title"] or "Untitled"
#         meta = format_meta_line(r)
#         sentences = top_sentences_for_record(query, r, max_sents=max_sents_per_rec)
#         if not sentences:
#             sentences = [{"sentence": (r.get("abstract") or "")[:800] + "...", "label": inline_citation_label(r), "href": best_href(r) or ""}]
#         # ReportLab hyperlink via <link href="...">...</link>
#         body_parts = []
#         for item in sentences:
#             if item["href"]:
#                 label_html = f'<link href="{html_escape(item["href"])}">{html_escape(item["label"])}</link>'
#             else:
#                 label_html = html_escape(item["label"])
#             body_parts.append(f'{html_escape(item["sentence"])} {label_html}')
#         body_html = " ".join(body_parts)
#
#         add_len = len(title) + len(meta) + len(body_html) + 40
#         if used + add_len > char_budget:
#             break
#
#         flow.append(Paragraph(html_escape(title), styles["PaperTitle"]))
#         if meta:
#             flow.append(Paragraph(html_escape(meta), styles["Meta"]))
#         flow.append(Paragraph(body_html, styles["Body"]))
#         used += add_len
#
#     return flow
#
# def build_pdf_list_exact_pages(query: str, records: List[Dict[str, Any]], one_image: Optional[bytes], target_pages: int) -> bytes:
#     if not PYPDF_AVAILABLE:
#         logger.warning("pypdf not installed; exact page control disabled. Install pypdf for exact page count.")
#         buf = io.BytesIO()
#         doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
#         flow = make_flow_for_budget(query, records, doc_width=doc.width, char_budget=target_pages * CHAR_BUDGET_PER_PAGE, target_pages=target_pages)
#         if one_image:
#             try:
#                 flow.insert(1, _img_flowable_from_bytes(one_image, max_width=doc.width))
#                 flow.insert(2, Spacer(1, 10))
#             except Exception:
#                 pass
#         doc.build(flow)
#         return buf.getvalue()
#
#     # Iteratively adjust char budget to hit exact pages
#     max_iters = 10
#     budget = target_pages * CHAR_BUDGET_PER_PAGE
#     last_pdf = b""
#     for i in range(max_iters):
#         buf = io.BytesIO()
#         doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
#         flow = make_flow_for_budget(query, records, doc_width=doc.width, char_budget=int(budget), target_pages=target_pages)
#         if one_image:
#             try:
#                 flow.insert(1, _img_flowable_from_bytes(one_image, max_width=doc.width))
#                 flow.insert(2, Spacer(1, 10))
#             except Exception:
#                 pass
#         doc.build(flow)
#         pdf_bytes = buf.getvalue()
#         last_pdf = pdf_bytes
#         buf.close()
#         try:
#             reader = PdfReader(io.BytesIO(pdf_bytes))
#             pages = len(reader.pages)
#         except Exception:
#             pages = 0
#         logger.info("PDF iteration %s: pages=%s target=%s budget=%s", i + 1, pages, target_pages, int(budget))
#         if pages == target_pages:
#             return pdf_bytes
#         if pages > target_pages:
#             budget *= 0.88
#             continue
#         # pages < target: increase budget
#         budget *= 1.18
#         if budget > (target_pages + 2) * CHAR_BUDGET_PER_PAGE:
#             budget = (target_pages + 2) * CHAR_BUDGET_PER_PAGE
#
#     # If still under target, pad with PageBreaks to reach exact page count
#     try:
#         base_reader = PdfReader(io.BytesIO(last_pdf))
#         current_pages = len(base_reader.pages)
#     except Exception:
#         current_pages = 0
#
#     if current_pages < target_pages:
#         buf = io.BytesIO()
#         doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
#         flow = make_flow_for_budget(query, records, doc_width=doc.width, char_budget=(target_pages * CHAR_BUDGET_PER_PAGE * 2), target_pages=max(target_pages, 20))
#         if one_image:
#             try:
#                 flow.insert(1, _img_flowable_from_bytes(one_image, max_width=doc.width))
#                 flow.insert(2, Spacer(1, 10))
#             except Exception:
#                 pass
#         for _ in range(target_pages - current_pages):
#             flow.append(PageBreak())
#             flow.append(Paragraph(" ", getSampleStyleSheet()["Normal"]))
#         doc.build(flow)
#         return buf.getvalue()
#     return last_pdf
#
# # DOCX builder (cannot guarantee exact pages due to Word rendering)
# def add_hyperlink(paragraph, url, text):
#     part = paragraph.part
#     r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
#     hyperlink = OxmlElement('w:hyperlink')
#     hyperlink.set(qn('r:id'), r_id)
#     new_run = OxmlElement('w:r')
#     rPr = OxmlElement('w:rPr')
#     u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
#     color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); rPr.append(color)
#     new_run.append(rPr)
#     t = OxmlElement('w:t'); t.text = text; new_run.append(t)
#     hyperlink.append(new_run)
#     paragraph._p.append(hyperlink)
#     return paragraph
#
# def build_docx_list(query: str, records: List[Dict[str, Any]], one_image: Optional[bytes], max_pages: int) -> bytes:
#     if not DOCX_AVAILABLE:
#         raise RuntimeError("DOCX output requested but python-docx is not installed.")
#     doc = DocxDocument()
#     doc.styles["Normal"].font.name = "Calibri"
#     doc.styles["Normal"].font.size = Pt(10)
#
#     h = doc.add_paragraph()
#     t = h.add_run(f"{APP_TITLE}: {query}")
#     t.bold = True
#     t.font.size = Pt(16)
#
#     if one_image:
#         try:
#             img_stream = io.BytesIO(one_image)
#             doc.add_picture(img_stream, width=Inches(6.0))
#         except Exception:
#             pass
#
#     char_budget = max_pages * CHAR_BUDGET_PER_PAGE
#     used = 0
#     max_sents_per_rec = dynamic_max_sents(max_pages)
#
#     for r in records:
#         title_p = doc.add_paragraph()
#         tr = title_p.add_run(r["title"] or "Untitled")
#         tr.bold = True
#         tr.font.size = Pt(12)
#
#         meta = format_meta_line(r)
#         if meta:
#             mp = doc.add_paragraph()
#             mr = mp.add_run(meta)
#             mr.italic = True
#
#         sentences = top_sentences_for_record(query, r, max_sents=max_sents_per_rec)
#         if not sentences:
#             sentences = [{"sentence": (r.get("abstract") or "")[:800] + "...", "label": inline_citation_label(r), "href": best_href(r) or ""}]
#         body_p = doc.add_paragraph()
#         for idx, item in enumerate(sentences):
#             body_p.add_run(item["sentence"] + " ")
#             if item["href"]:
#                 add_hyperlink(body_p, item["href"], item["label"])
#             else:
#                 body_p.add_run(item["label"])
#             if idx != len(sentences) - 1:
#                 body_p.add_run(" ")
#
#         add_len = len(r["title"] or "") + len(meta) + sum(len(x["sentence"]) + len(x["label"]) for x in sentences) + 40
#         if used + add_len > char_budget:
#             break
#         used += add_len
#
#     out = io.BytesIO()
#     doc.save(out)
#     return out.getvalue()
#
# # ------------------------------------------------------------------------------
# # Endpoints
# # ------------------------------------------------------------------------------
#
# @app.route("/report")
# def report():
#     if LOCAL_ROOT_DIR and os.path.isdir(LOCAL_ROOT_DIR):
#         cleanup_dir(LOCAL_ROOT_DIR)
#     temp_dir: Optional[str] = None
#     try:
#         raw_query = (request.args.get("query") or "").strip()
#         if not raw_query:
#             return jsonify({"error": "Missing query"}), 400
#
#         # Parse params
#         try:
#             max_pages = int(request.args.get("max_pages", DEFAULT_MAX_PAGES))
#             max_pages = max(1, min(max_pages, 100))  # up to 100 pages
#         except Exception:
#             max_pages = DEFAULT_MAX_PAGES
#
#         file_type = (request.args.get("file", "pdf") or "pdf").lower()
#         try:
#             max_results = int(request.args.get("max_results", DEFAULT_MAX_RESULTS))
#             max_results = max(10, min(max_results, 500))
#         except Exception:
#             max_results = DEFAULT_MAX_RESULTS
#
#         try:
#             since = int(request.args.get("since", DEFAULT_SINCE))
#         except Exception:
#             since = DEFAULT_SINCE
#
#         sources_param = request.args.get("sources", ",".join(sorted(AUTHORIZED_SOURCES)))
#         sources = [s.strip().lower() for s in sources_param.split(",") if s.strip()]
#         sources = [s for s in sources if s in AUTHORIZED_SOURCES]
#         if not sources:
#             sources = list(AUTHORIZED_SOURCES)
#
#         # Archive override via query param (default False to favor temp cleanup)
#         archive_override = request.args.get("archive")
#         archive = ARCHIVE_PDFS_DEFAULT if archive_override is None else (archive_override.strip().lower() in ("1", "true", "yes"))
#
#         # Choose storage dir: temp per request (default) or persistent archive folder
#         if archive:
#             storage_dir = get_archive_dir(raw_query)
#         else:
#             temp_dir = make_request_temp_dir()
#             storage_dir = temp_dir
#
#         logger.info('Report request: query=%r, sources=%s, max_pages=%s, file=%s, max_results=%s, since=%s, archive=%s, storage=%s',
#                     raw_query, sources, max_pages, file_type, max_results, since, archive, storage_dir)
#
#         # Fetch from providers using query variants (typo tolerance)
#         variants = query_variants(raw_query)
#         per_provider = max(5, min(100, math.ceil(max_results / max(1, len(sources)))))
#         all_records: List[Dict[str, Any]] = []
#         for vq in variants:
#             for src in sources:
#                 try:
#                     recs = FETCHERS[src](vq, per_provider, since)
#                     all_records.extend(recs)
#                 except Exception as e:
#                     logger.warning("Fetcher %s failed for variant %r: %s", src, vq, e)
#
#         if not all_records:
#             return jsonify({"error": "No results found from authorized sources with abstracts. Try a different query."}), 404
#
#         # Dedupe, prefer OA, rerank (downloads some PDFs into storage_dir for snippet extraction)
#         records = dedupe_records(all_records)
#         records = sorted(records, key=lambda r: 0 if r.get("pdf_url") else 1)  # OA first
#         ranked = rerank_records(raw_query, records, dest_dir=storage_dir, use_pdf_snippets=True, snippet_limit=8)
#
#         # For image extraction, we may download one or two more PDFs to storage_dir
#         one_image = extract_one_image_from_records(ranked, dest_dir=storage_dir)
#
#         # Persist manifest only when archiving (storage_dir is ./local/slug)
#         if archive:
#             # Ensure PDFs for all ranked with pdf_url are saved into archive folder
#             for r in ranked:
#                 if r.get("pdf_url"):
#                     ensure_local_pdf(r, storage_dir)
#             write_manifest(storage_dir, raw_query, ranked)
#
#         # Build output in-memory (no file saved on disk for the final PDF/DOCX)
#         if file_type == "pdf":
#             out_bytes = build_pdf_list_exact_pages(raw_query, ranked, one_image, target_pages=max_pages)
#          #   fname = f"report_{re.sub(r'\W+', '_', raw_query)}.pdf"
#          #    fname = f"report_{re.sub(r'\W+', '_', raw_query)}.pdf"
#             fname = f'report_{re.sub(r"W+", "_", raw_query)}.pdf'
#             resp = send_file(io.BytesIO(out_bytes), mimetype="application/pdf", as_attachment=True, download_name=fname)
#         elif file_type == "docx":
#             out_bytes = build_docx_list(raw_query, ranked, one_image, max_pages=max_pages)
#             #fname = f"report_{re.sub(r'\W+', '_', raw_query)}.docx"
#             # fname = f"report_{re.sub(r'\W+', '_', raw_query)}.docx"
#             fname = f"report_{re.sub(r'W+', '_', raw_query)}.docx"
#             resp = send_file(io.BytesIO(out_bytes), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name=fname)
#         else:
#             return jsonify({"error": "Unsupported file type. Use file=pdf or file=docx."}), 400
#
#         return resp
#     finally:
#         # If we used a per-request temp dir (archive=0), clean it up now.
#         # At this point the response bytes have been created in-memory and sent back.
#         if temp_dir and os.path.isdir(LOCAL_ROOT_DIR):
#             cleanup_dir(LOCAL_ROOT_DIR)
#         if os.path.isdir(LOCAL_ROOT_DIR):
#             cleanup_dir(LOCAL_ROOT_DIR)
#
# # ------------------------------------------------------------------------------
# # Main
# # ------------------------------------------------------------------------------
#
# if __name__ == "__main__":
#     desired_debug = os.environ.get("ALLOW_DEBUG", "").lower() in ("1", "true", "yes")
#     opts = flask_run_options(desired_debug)
#     port = int(os.environ.get("PORT", "8002"))
#     logger.info("Starting server on 0.0.0.0:%s (debug=%s)", port, opts.get("debug"))
#     app.run(host="0.0.0.0", port=port, **opts)



# """
# Literature Snapshot API — Multi-source, OA-first, typo-tolerant search
# List format with per-paper summaries and inline clickable citations
# Exact PDF page count control (supports large values, e.g., 50+ pages)
# Temp folder lifecycle for all downloaded files (auto-cleanup when not archived)
#
# Key additions in this version
# - All transient downloads (e.g., source PDFs used for snippet/image extraction) go into a per-request temp folder
#   under TEMP_ROOT_DIR (defaults to OS temp, e.g., /tmp/literature_snapshot or %TEMP%\literature_snapshot).
# - If archive=0 (default behavior you want), the per-request temp folder is deleted at the end of the request,
#   so nothing remains on disk once the PDF/DOCX has been generated and streamed to the user.
# - If archive=1, files are saved under ./local/{query_slug}/ and NOT deleted (plus manifest.json is written).
#
# Endpoint
# - GET /report?query=...&max_pages=50&file=pdf
#   Parameters:
#     - query       (required): topic (typos tolerated)
#     - max_pages   (required for exact pages): integer 1–100 (supports 50+)
#     - file        (optional): pdf (default) or docx
#     - max_results (optional): 10–500 (default 120)
#     - sources     (optional): comma-separated in [openalex,arxiv,eupmc,crossref,s2] (default all)
#     - since       (optional): min year (default 2015)
#     - archive     (optional): 1 or 0 (default 0 here) — when 0, temp files are cleaned up after response
#
# Notes
# - For exact page counts, ensure pypdf is installed (see requirements.txt).
# - Inline citations are clickable links that open the source page/PDF in a new tab.
# """
#
# from __future__ import annotations
#
import io
import os
import re
import math
import html
import json
import time
import uuid
import shutil
import logging
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

import requests
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS  # CORS for browser fetches

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import utils

# PDF page counting (to guarantee exact pages)
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except Exception:
    PYPDF_AVAILABLE = False

# Optional DOCX output (exact pages not guaranteed in DOCX)
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Optional PyMuPDF for image/snippet extraction
try:
    import fitz  # type: ignore
    PYMUPDF_AVAILABLE = True
except Exception:
    PYMUPDF_AVAILABLE = False

# Optional scikit-learn for TF-IDF ranking
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

# ------------------------------------------------------------------------------
# Config and constants
# ------------------------------------------------------------------------------

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

APP_TITLE = "Literature Snapshot"

DEFAULT_MAX_PAGES = 2
DEFAULT_MAX_RESULTS = 120   # higher default to support large page counts
DEFAULT_SINCE = 2015
AUTHORIZED_SOURCES = {"openalex", "arxiv", "eupmc", "crossref", "s2"}

# External APIs
OPENALEX_BASE = "https://api.openalex.org/works"
ARXIV_BASE = "http://export.arxiv.org/api/query"
EUROPE_PMC_BASE = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
CROSSREF_BASE = "https://api.crossref.org/works"
S2_BASE = "https://api.semanticscholar.org/graph/v1/paper/search"

# Character budget heuristic per page (tuned for list format and 10pt body)
CHAR_BUDGET_PER_PAGE = 2200

# Representative image filter
MIN_IMG_WIDTH = 300
MIN_IMG_HEIGHT = 150
MIN_IMG_SIZE_BYTES = 20 * 1024

# Archive root for persisted PDFs and manifests (when archive=1)
LOCAL_ROOT_DIR = "./local"

# Temp root for per-request transient files
TEMP_ROOT_DIR = os.environ.get("TEMP_ROOT_DIR") or os.path.join(tempfile.gettempdir(), "literature_snapshot")

# Default archiving toggle for this build: do NOT keep files unless explicitly requested
ARCHIVE_PDFS_DEFAULT = False

# ------------------------------------------------------------------------------
# App bootstrap (CORS)
# ------------------------------------------------------------------------------

app = Flask(__name__)
CORS(
    app,
    resources={r"/*": {"origins": "*"}},
    expose_headers=["Content-Disposition"]
)
os.makedirs(LOCAL_ROOT_DIR, exist_ok=True)
os.makedirs(TEMP_ROOT_DIR, exist_ok=True)

def detect_multiprocessing_availability() -> bool:
    try:
        import _multiprocessing  # noqa: F401
        return True
    except Exception:
        return False

def flask_run_options(desired_debug: bool) -> Dict[str, Any]:
    if desired_debug and detect_multiprocessing_availability():
        return {"debug": True, "use_reloader": True}
    else:
        return {"debug": False, "use_reloader": False}

# ------------------------------------------------------------------------------
# Utilities
# ------------------------------------------------------------------------------

def normspace(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()

def strip_html(jats_or_html: str) -> str:
    no_tags = re.sub(r"<[^>]+>", " ", jats_or_html or "")
    return normspace(html.unescape(no_tags))

def sentence_case(title: str) -> str:
    if not title:
        return title
    t = title.strip()
    return t[:1].upper() + t[1:]

def year_from_date(d: Optional[str]) -> Optional[int]:
    if not d:
        return None
    m = re.match(r"(\d{4})", d)
    return int(m.group(1)) if m else None

def html_escape(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def safe_request_json(url: str, params: Dict[str, Any], timeout: int = 30, headers: Optional[Dict[str, str]] = None) -> Optional[Dict[str, Any]]:
    try:
        r = requests.get(url, params=params, timeout=timeout, headers=headers or {})
        r.raise_for_status()
        return r.json()
    except Exception as e:
        logger.warning("Request failed: %s params=%s error=%s", url, params, e)
        return None

def safe_slug(s: str) -> str:
    s = normspace(s).lower()
    s = re.sub(r"[^\w\-]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s[:80] or f"query_{int(time.time())}"

def safe_filename(s: str) -> str:
    s = normspace(s)
    s = re.sub(r"[^\w\-]+", "_", s)
    return s[:120] or f"paper_{int(time.time())}"

# ------------------------------------------------------------------------------
# Temp and archive dirs
# ------------------------------------------------------------------------------

def make_request_temp_dir() -> str:
    """
    Create a unique temp folder for this request under TEMP_ROOT_DIR.
    """
    uid = uuid.uuid4().hex[:12]
    path = os.path.join(TEMP_ROOT_DIR, f"req_{uid}")
    os.makedirs(path, exist_ok=True)
    return path

def cleanup_dir(path: Optional[str]) -> None:
    if not path:
        return
    try:
        shutil.rmtree(path, ignore_errors=True)
        logger.info("Cleaned temp folder: %s", path)
    except Exception as e:
        logger.warning("Failed to clean temp folder %s: %s", path, e)

def get_archive_dir(query: str) -> str:
    slug = safe_slug(query)
    path = os.path.join(LOCAL_ROOT_DIR, slug)
    os.makedirs(path, exist_ok=True)
    return path

# ------------------------------------------------------------------------------
# Query normalization for typo tolerance
# ------------------------------------------------------------------------------

def query_variants(q: str) -> List[str]:
    """
    Return a list of query variants to tolerate simple spelling/formatting mistakes.
    """
    q = q or ""
    base = normspace(q)
    lower = base.lower()
    collapsed = re.sub(r"(.)\1{2,}", r"\1\1", lower)
    punct_trim = re.sub(r"[\"'`~^|]+", " ", lower)
    alnum = re.sub(r"[^\w\s]+", " ", lower)
    seen, out = set(), []
    for v in [base, lower, collapsed, punct_trim, alnum]:
        v = normspace(v)
        if v and v not in seen:
            seen.add(v); out.append(v)
    return out

# ------------------------------------------------------------------------------
# Record helpers and inline citations
# ------------------------------------------------------------------------------

def unify_record(
    source: str,
    title: Optional[str],
    abstract: Optional[str],
    year: Optional[int],
    url: Optional[str],
    pdf_url: Optional[str],
    doi: Optional[str],
    authors: Optional[List[str]] = None,
    venue: Optional[str] = None,
) -> Dict[str, Any]:
    return {
        "source": source,
        "title": normspace(title or ""),
        "abstract": normspace(abstract or ""),
        "year": year,
        "url": url,
        "pdf_url": pdf_url,
        "doi": (doi or "").lower() or None,
        "authors": authors or [],
        "venue": normspace(venue or ""),
        "_local_pdf": None,
        "_local_pdf_rel": None,
        "_snippet": None,
    }

def is_valid_record(rec: Dict[str, Any]) -> bool:
    if not rec["title"] or not rec["abstract"]:
        return False
    if rec["year"] and rec["year"] < DEFAULT_SINCE:
        return False
    return True

def dedupe_records(recs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    out = []
    for r in recs:
        key = r.get("doi") or r["title"].lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(r)
    return out

def last_names(authors: List[str]) -> List[str]:
    ln = []
    for a in authors or []:
        parts = [p for p in a.replace(",", " ").split() if p]
        if parts:
            ln.append(parts[-1])
    return ln

def inline_citation_label(rec: Dict[str, Any]) -> str:
    ln = last_names(rec.get("authors") or [])
    y = rec.get("year")
    ystr = str(y) if y else "n.d."
    if not ln:
        return f"({ystr})"
    if len(ln) == 1:
        return f"({ln[0]}, {ystr})"
    if len(ln) == 2:
        return f"({ln[0]} & {ln[1]}, {ystr})"
    return f"({ln[0]} et al., {ystr})"

def best_href(rec: Dict[str, Any]) -> Optional[str]:
    return rec.get("pdf_url") or rec.get("url")

# ------------------------------------------------------------------------------
# Providers
# ------------------------------------------------------------------------------

def reconstruct_openalex_abstract(inv_index: Optional[Dict[str, List[int]]]) -> Optional[str]:
    if not inv_index:
        return None
    max_pos = -1
    for positions in inv_index.values():
        if positions:
            max_pos = max(max_pos, max(positions))
    if max_pos < 0:
        return None
    tokens: List[Optional[str]] = [None] * (max_pos + 1)
    for token, positions in inv_index.items():
        for pos in positions:
            if 0 <= pos < len(tokens):
                tokens[pos] = token
    return normspace(" ".join(t or "" for t in tokens))

def fetch_openalex(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
    params = {
        "filter": f"concepts.display_name.search:{query},display_name.search:{query},has_fulltext:true,has_abstract:true,from_publication_date:{since}-01-01",
        "sort": "relevance_score:desc,cited_by_count:desc",
        "per-page": min(limit, 25),
        "page": 1,
    }
    data = safe_request_json(OPENALEX_BASE, params)
    results = data.get("results", []) if data else []
    out: List[Dict[str, Any]] = []
    for w in results:
        abstract = w.get("abstract")
        if not abstract and w.get("abstract_inverted_index"):
            abstract = reconstruct_openalex_abstract(w["abstract_inverted_index"])
        title = w.get("title")
        year = w.get("publication_year")
        doi = w.get("doi")
        venue = (w.get("host_venue") or {}).get("display_name")
        authors = []
        for a in (w.get("authorships") or []):
            nm = (a.get("author") or {}).get("display_name")
            if nm: authors.append(nm)
        pdf_url = None
        try:
            best = w.get("best_oa_location") or {}
            if best.get("pdf_url"): pdf_url = best["pdf_url"]
        except Exception:
            pass
        if not pdf_url:
            for loc in (w.get("oa_locations") or []):
                if loc.get("pdf_url"): pdf_url = loc["pdf_url"]; break
        landing = (w.get("primary_location") or {}).get("landing_page_url")
        url = landing or (w.get("host_venue") or {}).get("url") or w.get("id")
        rec = unify_record("OPENALEX", title, abstract, year, url, pdf_url, doi, authors, venue)
        if is_valid_record(rec): out.append(rec)
        if len(out) >= limit: break
    return out

def fetch_arxiv(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
    per_page = min(limit, 25)
    params = {
        "search_query": f'all:{query}',
        "start": 0,
        "max_results": per_page,
        "sortBy": "relevance",
        "sortOrder": "descending",
    }
    try:
        r = requests.get(ARXIV_BASE, params=params, timeout=30)
        r.raise_for_status()
        import xml.etree.ElementTree as ET
        root = ET.fromstring(r.text)
        ns = {"a": "http://www.w3.org/2005/Atom"}
        out: List[Dict[str, Any]] = []
        for entry in root.findall("a:entry", ns):
            title = normspace(entry.findtext("a:title", default="", namespaces=ns))
            abstract = normspace(entry.findtext("a:summary", default="", namespaces=ns))
            published = entry.findtext("a:published", default="", namespaces=ns)
            year = year_from_date(published)
            if year and year < since: continue
            url = entry.findtext("a:id", default="", namespaces=ns)
            pdf_url = None
            for link in entry.findall("a:link", ns):
                if link.attrib.get("type") == "application/pdf":
                    pdf_url = link.attrib.get("href"); break
            authors = [normspace(a.findtext("a:name", default="", namespaces=ns)) for a in entry.findall("a:author", ns)]
            venue = "arXiv"
            rec = unify_record("ARXIV", title, abstract, year, url, pdf_url, None, authors, venue)
            if is_valid_record(rec): out.append(rec)
            if len(out) >= limit: break
        return out
    except Exception as e:
        logger.warning("arXiv fetch failed: %s", e)
        return []

def fetch_europe_pmc(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
    params = {
        "query": f'({query}) AND PUB_YEAR:[{since} TO 3000] AND HAS_ABSTRACT:Y',
        "resultType": "core",
        "pageSize": min(limit, 25),
        "format": "json",
    }
    data = safe_request_json(EUROPE_PMC_BASE, params)
    results = data.get("resultList", {}).get("result", []) if data else []
    out: List[Dict[str, Any]] = []
    for it in results:
        title = it.get("title")
        abstract = it.get("abstractText")
        year = int(it["pubYear"]) if it.get("pubYear") and str(it["pubYear"]).isdigit() else None
        if year and year < since: continue
        doi = it.get("doi")
        url = None
        if it.get("fullTextUrlList") and it["fullTextUrlList"].get("fullTextUrl"):
            url = it["fullTextUrlList"]["fullTextUrl"][0].get("url")
        url = url or it.get("doi")
        pdf_url = None
        for ft in (it.get("fullTextUrlList", {}) or {}).get("fullTextUrl", []) or []:
            u = ft.get("url") or ""
            if ft.get("documentStyle") == "pdf" or ("pdf" in u.lower() and ft.get("availability", "").lower().startswith("open")):
                pdf_url = u; break
        venue = it.get("journalTitle") or "Europe PMC"
        authors: List[str] = []
        if it.get("authorList") and it["authorList"].get("author"):
            for a in it["authorList"]["author"]:
                nm = a.get("fullName") or a.get("lastName") or ""
                if a.get("firstName"): nm = f"{a['firstName']} {nm}".strip()
                if nm: authors.append(nm)
        elif it.get("authorString"):
            authors = [x.strip() for x in re.split(r"[;,]", it["authorString"]) if x.strip()]
        rec = unify_record("EUROPE PMC", title, abstract, year, url, pdf_url, doi, authors, venue)
        if is_valid_record(rec): out.append(rec)
        if len(out) >= limit: break
    return out

def fetch_crossref(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
    params = {
        "query.title": query,
        "filter": f"from-pub-date:{since}-01-01,type:journal-article,has-abstract:true",
        "rows": min(limit, 25),
        "sort": "relevance",
        "order": "desc",
    }
    try:
        r = requests.get(CROSSREF_BASE, params=params, timeout=30)
        r.raise_for_status()
        data = r.json()
        items = data.get("message", {}).get("items", []) or []
        out: List[Dict[str, Any]] = []
        for it in items:
            title = normspace(" ".join(it.get("title") or []))
            abstract_raw = it.get("abstract")
            abstract = strip_html(abstract_raw) if abstract_raw else ""
            year = None
            if it.get("issued", {}).get("date-parts"):
                year = it["issued"]["date-parts"][0][0]
            if year and year < since: continue
            doi = it.get("DOI")
            url = it.get("URL")
            pdf_url = None
            for link in it.get("link", []) or []:
                if link.get("content-type") == "application/pdf":
                    pdf_url = link.get("URL"); break
            venue = ""
            if it.get("container-title"):
                venue = " ".join(it.get("container-title"))
            authors: List[str] = []
            for a in it.get("author", []) or []:
                given = a.get("given", ""); family = a.get("family", "")
                nm = " ".join([given, family]).strip() or (a.get("name") or "")
                if nm: authors.append(nm)
            oa_license = any(("open" in (lic.get("URL", "") + lic.get("content-version", "")).lower()) for lic in it.get("license", []) or [])
            rec = unify_record("CROSSREF", title, abstract, year, url, pdf_url, doi, authors, venue)
            if is_valid_record(rec) and (pdf_url or oa_license): out.append(rec)
            if len(out) >= limit: break
        return out
    except Exception as e:
        logger.warning("Crossref fetch failed: %s", e)
        return []

def fetch_semantic_scholar(query: str, limit: int, since: int) -> List[Dict[str, Any]]:
    params = {
        "query": query,
        "fields": "title,abstract,year,url,openAccessPdf,externalIds,authors,venue",
        "limit": min(limit, 25),
        "offset": 0,
    }
    headers = {}
    if os.environ.get("S2_API_KEY"):
        headers["x-api-key"] = os.environ["S2_API_KEY"]
    data = safe_request_json(S2_BASE, params, headers=headers)
    if not data:
        return []
    out: List[Dict[str, Any]] = []
    for it in data.get("data", []) or []:
        year = it.get("year")
        if year and year < since: continue
        title = it.get("title") or ""
        abstract = it.get("abstract") or ""
        doi = (it.get("externalIds") or {}).get("DOI")
        url = it.get("url")
        pdf_url = (it.get("openAccessPdf") or {}).get("url")
        authors = [a.get("name") for a in (it.get("authors") or []) if a.get("name")]
        venue = it.get("venue") or ""
        rec = unify_record("SEMANTIC SCHOLAR", title, abstract, year, url, pdf_url, doi, authors, venue)
        if is_valid_record(rec): out.append(rec)
        if len(out) >= limit: break
    return out

FETCHERS = {
    "openalex": fetch_openalex,
    "arxiv": fetch_arxiv,
    "eupmc": fetch_europe_pmc,
    "crossref": fetch_crossref,
    "s2": fetch_semantic_scholar,
}

# ------------------------------------------------------------------------------
# Per-request storage handling (temp vs archive)
# ------------------------------------------------------------------------------

def download_pdf_to_dir(url: str, rec: Dict[str, Any], dest_dir: str) -> Optional[str]:
    if not url:
        return None
    try:
        r = requests.get(url, timeout=60, stream=True)
        r.raise_for_status()
        ts = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        tag = rec.get("doi") or rec.get("title") or "paper"
        filename = f"{rec.get('source','src')}_{safe_filename(tag)}_{ts}.pdf"
        local_path = os.path.join(dest_dir, filename)
        with open(local_path, "wb") as f:
            for chunk in r.iter_content(chunk_size=8192):
                if chunk: f.write(chunk)
        abs_path = os.path.abspath(local_path)
        rec["_local_pdf"] = abs_path
        rec["_local_pdf_rel"] = None
        return abs_path
    except Exception as e:
        logger.warning("Failed to download PDF %s: %s", url, e)
        return None

def ensure_local_pdf(rec: Dict[str, Any], dest_dir: str) -> Optional[str]:
    if rec.get("_local_pdf") and os.path.exists(rec["_local_pdf"]):
        return rec["_local_pdf"]
    if not rec.get("pdf_url"):
        return None
    return download_pdf_to_dir(rec["pdf_url"], rec, dest_dir)

def write_manifest(dest_dir: str, query: str, records: List[Dict[str, Any]]) -> str:
    payload = {
        "query": query,
        "generated_at_utc": datetime.utcnow().isoformat() + "Z",
        "folder": os.path.abspath(dest_dir),
        "records": []
    }
    for r in records:
        payload["records"].append({
            "source": r.get("source"),
            "title": r.get("title"),
            "authors": r.get("authors"),
            "year": r.get("year"),
            "venue": r.get("venue"),
            "doi": r.get("doi"),
            "url": r.get("url"),
            "pdf_url": r.get("pdf_url"),
            "local_pdf": r.get("_local_pdf"),  # absolute path in archive dir
        })
    manifest_path = os.path.join(dest_dir, "manifest.json")
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    logger.info("Wrote manifest for query: %s", os.path.basename(dest_dir))
    return manifest_path

# ------------------------------------------------------------------------------
# Ranking, snippets, and list-style summaries
# ------------------------------------------------------------------------------

def compute_similarity_scores(query: str, texts: List[str]) -> List[float]:
    if SKLEARN_AVAILABLE:
        vec = TfidfVectorizer(stop_words="english", max_features=20000)
        X = vec.fit_transform([query] + texts)
        sims = cosine_similarity(X[0:1], X[1:]).flatten().tolist()
        return sims
    # Fallback: token overlap
    q_tokens = set(re.findall(r"\w+", query.lower()))
    sims: List[float] = []
    for t in texts:
        t_tokens = set(re.findall(r"\w+", (t or "").lower()))
        inter = len(q_tokens & t_tokens)
        denom = math.sqrt(max(len(q_tokens), 1) * max(len(t_tokens), 1))
        sims.append(inter / denom if denom else 0.0)
    return sims

def extract_text_snippet_from_pdf(local_pdf_path: str, max_chars: int = 1400) -> str:
    if not PYMUPDF_AVAILABLE or not local_pdf_path:
        return ""
    doc = None
    try:
        doc = fitz.open(local_pdf_path)
        txts: List[str] = []
        for i in range(min(len(doc), 3)):
            page = doc.load_page(i)
            txts.append(page.get_text("text"))
        snippet = normspace(" ".join(txts))[:max_chars]
        return snippet
    except Exception:
        return ""
    finally:
        if doc:
            doc.close()

def rerank_records(query: str, records: List[Dict[str, Any]], dest_dir: str, use_pdf_snippets: bool = True, snippet_limit: int = 8) -> List[Dict[str, Any]]:
    enriched: List[Dict[str, Any]] = []
    for r in records[:snippet_limit]:
        if use_pdf_snippets and r.get("pdf_url"):
            local_pdf = ensure_local_pdf(r, dest_dir)
            if local_pdf:
                snip = extract_text_snippet_from_pdf(local_pdf, max_chars=1400)
                if snip:
                    r = dict(r)
                    r["_snippet"] = snip
        enriched.append(r)
    enriched.extend(records[snippet_limit:])
    texts = [f"{rr['title']} {rr['abstract']} {' ' + (rr['_snippet'] or '') if rr.get('_snippet') else ''}" for rr in enriched]
    sims = compute_similarity_scores(query, texts)
    ranked = [rec for _, rec in sorted(zip(sims, enriched), key=lambda x: x[0], reverse=True)]
    return ranked

def split_sentences(text: str) -> List[str]:
    s = normspace(text)
    if not s:
        return []
    parts = re.split(r'(?<=[.!?])\s+', s)
    return [p.strip() for p in parts if p.strip()]

def top_sentences_for_record(query: str, rec: Dict[str, Any], max_sents: int = 3) -> List[Dict[str, str]]:
    base = rec.get("abstract") or ""
    if rec.get("_snippet"):
        base = base + " " + rec["_snippet"]
    sents = split_sentences(base)
    if not sents:
        return []
    # Score by token overlap to pick the most relevant sentences
    q_tokens = set(re.findall(r"\w+", query.lower()))
    scored = []
    for s in sents:
        tokens = set(re.findall(r"\w+", s.lower()))
        inter = len(q_tokens & tokens)
        denom = math.sqrt(max(len(q_tokens), 1) * max(len(tokens), 1))
        score = inter / denom if denom else 0.0
        scored.append((score, s))
    scored.sort(key=lambda x: x[0], reverse=True)
    chosen = [s for _, s in scored[:max_sents]]
    label = inline_citation_label(rec)
    href = best_href(rec)
    out = []
    for s in chosen:
        s_clean = s.rstrip()
        if not s_clean.endswith(('.', '!', '?')):
            s_clean += '.'
        out.append({"sentence": s_clean, "label": label, "href": href or ""})
    return out

def format_meta_line(rec: Dict[str, Any]) -> str:
    parts = []
    if rec.get("source"):
        parts.append(rec["source"])
    if rec.get("year"):
        parts.append(str(rec["year"]))
    if rec.get("venue"):
        parts.append(rec["venue"])
    return " • ".join(parts)

# ------------------------------------------------------------------------------
# Image extraction
# ------------------------------------------------------------------------------

def _is_good_image(pix: "fitz.Pixmap", img_bytes: bytes) -> bool:
    return (pix.width >= MIN_IMG_WIDTH and pix.height >= MIN_IMG_HEIGHT and len(img_bytes) >= MIN_IMG_SIZE_BYTES)

def extract_one_image_from_records(records: List[Dict[str, Any]], dest_dir: str) -> Optional[bytes]:
    if not PYMUPDF_AVAILABLE:
        return None
    for r in records:
        local_pdf = ensure_local_pdf(r, dest_dir)
        if not local_pdf:
            continue
        doc = None
        try:
            doc = fitz.open(local_pdf)
            for p in range(min(len(doc), 5)):
                page = doc.load_page(p)
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    img_bytes = base_image.get("image")
                    if not img_bytes:
                        continue
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n >= 5:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    if _is_good_image(pix, img_bytes):
                        return img_bytes
        except Exception as e:
            logger.warning("Image extraction failed for %s: %s", r.get("title"), e)
        finally:
            if doc:
                doc.close()
    return None

# ------------------------------------------------------------------------------
# Output generation — LIST FORMAT, exact pages for PDF
# ------------------------------------------------------------------------------

def _img_flowable_from_bytes(img_bytes: bytes, max_width: float) -> Image:
    bio = io.BytesIO(img_bytes)
    ir = utils.ImageReader(bio)
    iw, ih = ir.getSize()
    aspect = ih / iw
    width = min(max_width, iw)
    height = width * aspect
    bio.seek(0)
    return Image(bio, width=width, height=height)

def dynamic_max_sents(target_pages: int) -> int:
    base = 3
    extra = min(9, max(0, (target_pages - 5) // 5 * 2))
    return base + extra

def make_flow_for_budget(query: str, records: List[Dict[str, Any]], doc_width: float, char_budget: int, target_pages: int) -> List[Any]:
    styles = getSampleStyleSheet()
    if "TitleSmall" not in styles:
        styles.add(ParagraphStyle(name="TitleSmall", parent=styles['h2'], fontSize=16, leading=20, spaceAfter=12))
        styles.add(ParagraphStyle(name="PaperTitle", parent=styles['h4'], fontSize=12, leading=15, spaceAfter=2))
        styles.add(ParagraphStyle(name="Meta", parent=styles['Normal'], fontSize=9.5, leading=12, textColor="#555555", spaceAfter=4))
        styles.add(ParagraphStyle(name="Body", parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=10))
    flow: List[Any] = []
    flow.append(Paragraph(f"{APP_TITLE}: {html_escape(query)}", styles["TitleSmall"]))

    used = 0
    max_sents_per_rec = dynamic_max_sents(target_pages)

    for r in records:
        title = r["title"] or "Untitled"
        meta = format_meta_line(r)
        sentences = top_sentences_for_record(query, r, max_sents=max_sents_per_rec)
        if not sentences:
            sentences = [{"sentence": (r.get("abstract") or "")[:800] + "...", "label": inline_citation_label(r), "href": best_href(r) or ""}]
        # ReportLab hyperlink via <link href="...">...</link>
        body_parts = []
        for item in sentences:
            if item["href"]:
                label_html = f'<link href="{html_escape(item["href"])}">{html_escape(item["label"])}</link>'
            else:
                label_html = html_escape(item["label"])
            body_parts.append(f'{html_escape(item["sentence"])} {label_html}')
        body_html = " ".join(body_parts)

        add_len = len(title) + len(meta) + len(body_html) + 40
        if used + add_len > char_budget:
            break

        flow.append(Paragraph(html_escape(title), styles["PaperTitle"]))
        if meta:
            flow.append(Paragraph(html_escape(meta), styles["Meta"]))
        flow.append(Paragraph(body_html, styles["Body"]))
        used += add_len

    return flow

def build_pdf_list_exact_pages(query: str, records: List[Dict[str, Any]], one_image: Optional[bytes], target_pages: int) -> bytes:
    if not PYPDF_AVAILABLE:
        logger.warning("pypdf not installed; exact page control disabled. Install pypdf for exact page count.")
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        flow = make_flow_for_budget(query, records, doc_width=doc.width, char_budget=target_pages * CHAR_BUDGET_PER_PAGE, target_pages=target_pages)
        if one_image:
            try:
                flow.insert(1, _img_flowable_from_bytes(one_image, max_width=doc.width))
                flow.insert(2, Spacer(1, 10))
            except Exception:
                pass
        doc.build(flow)
        return buf.getvalue()

    # Iteratively adjust char budget to hit exact pages
    max_iters = 10
    budget = target_pages * CHAR_BUDGET_PER_PAGE
    last_pdf = b""
    for i in range(max_iters):
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        flow = make_flow_for_budget(query, records, doc_width=doc.width, char_budget=int(budget), target_pages=target_pages)
        if one_image:
            try:
                flow.insert(1, _img_flowable_from_bytes(one_image, max_width=doc.width))
                flow.insert(2, Spacer(1, 10))
            except Exception:
                pass
        doc.build(flow)
        pdf_bytes = buf.getvalue()
        last_pdf = pdf_bytes
        buf.close()
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            pages = len(reader.pages)
        except Exception:
            pages = 0
        logger.info("PDF iteration %s: pages=%s target=%s budget=%s", i + 1, pages, target_pages, int(budget))
        if pages == target_pages:
            return pdf_bytes
        if pages > target_pages:
            budget *= 0.88
            continue
        # pages < target: increase budget
        budget *= 1.18
        if budget > (target_pages + 2) * CHAR_BUDGET_PER_PAGE:
            budget = (target_pages + 2) * CHAR_BUDGET_PER_PAGE

    # If still under target, pad with PageBreaks to reach exact page count
    try:
        base_reader = PdfReader(io.BytesIO(last_pdf))
        current_pages = len(base_reader.pages)
    except Exception:
        current_pages = 0

    if current_pages < target_pages:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        flow = make_flow_for_budget(query, records, doc_width=doc.width, char_budget=(target_pages * CHAR_BUDGET_PER_PAGE * 2), target_pages=max(target_pages, 20))
        if one_image:
            try:
                flow.insert(1, _img_flowable_from_bytes(one_image, max_width=doc.width))
                flow.insert(2, Spacer(1, 10))
            except Exception:
                pass
        for _ in range(target_pages - current_pages):
            flow.append(PageBreak())
            flow.append(Paragraph(" ", getSampleStyleSheet()["Normal"]))
        doc.build(flow)
        return buf.getvalue()
    return last_pdf

# DOCX builder (cannot guarantee exact pages due to Word rendering)
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def build_docx_list(query: str, records: List[Dict[str, Any]], one_image: Optional[bytes], max_pages: int) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("DOCX output requested but python-docx is not installed.")
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    h = doc.add_paragraph()
    t = h.add_run(f"{APP_TITLE}: {query}")
    t.bold = True
    t.font.size = Pt(16)

    if one_image:
        try:
            img_stream = io.BytesIO(one_image)
            doc.add_picture(img_stream, width=Inches(6.0))
        except Exception:
            pass

    char_budget = max_pages * CHAR_BUDGET_PER_PAGE
    used = 0
    max_sents_per_rec = dynamic_max_sents(max_pages)

    for r in records:
        title_p = doc.add_paragraph()
        tr = title_p.add_run(r["title"] or "Untitled")
        tr.bold = True
        tr.font.size = Pt(12)

        meta = format_meta_line(r)
        if meta:
            mp = doc.add_paragraph()
            mr = mp.add_run(meta)
            mr.italic = True

        sentences = top_sentences_for_record(query, r, max_sents=max_sents_per_rec)
        if not sentences:
            sentences = [{"sentence": (r.get("abstract") or "")[:800] + "...", "label": inline_citation_label(r), "href": best_href(r) or ""}]
        body_p = doc.add_paragraph()
        for idx, item in enumerate(sentences):
            body_p.add_run(item["sentence"] + " ")
            if item["href"]:
                add_hyperlink(body_p, item["href"], item["label"])
            else:
                body_p.add_run(item["label"])
            if idx != len(sentences) - 1:
                body_p.add_run(" ")

        add_len = len(r["title"] or "") + len(meta) + sum(len(x["sentence"]) + len(x["label"]) for x in sentences) + 40
        if used + add_len > char_budget:
            break
        used += add_len

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# ------------------------------------------------------------------------------
# Endpoints
# ------------------------------------------------------------------------------

@app.route("/report")
def report():
    temp_dir: Optional[str] = None
    try:
        raw_query = (request.args.get("query") or "").strip()
        if not raw_query:
            return jsonify({"error": "Missing query"}), 400

        # Parse params
        try:
            max_pages = int(request.args.get("max_pages", DEFAULT_MAX_PAGES))
            max_pages = max(1, min(max_pages, 100))  # up to 100 pages
        except Exception:
            max_pages = DEFAULT_MAX_PAGES

        file_type = (request.args.get("file", "pdf") or "pdf").lower()
        try:
            max_results = int(request.args.get("max_results", DEFAULT_MAX_RESULTS))
            max_results = max(10, min(max_results, 500))
        except Exception:
            max_results = DEFAULT_MAX_RESULTS

        try:
            since = int(request.args.get("since", DEFAULT_SINCE))
        except Exception:
            since = DEFAULT_SINCE

        sources_param = request.args.get("sources", ",".join(sorted(AUTHORIZED_SOURCES)))
        sources = [s.strip().lower() for s in sources_param.split(",") if s.strip()]
        sources = [s for s in sources if s in AUTHORIZED_SOURCES]
        if not sources:
            sources = list(AUTHORIZED_SOURCES)

        # Archive override via query param (default False to favor temp cleanup)
        archive_override = request.args.get("archive")
        archive = ARCHIVE_PDFS_DEFAULT if archive_override is None else (archive_override.strip().lower() in ("1", "true", "yes"))

        # Choose storage dir: temp per request (default) or persistent archive folder
        if archive:
            storage_dir = get_archive_dir(raw_query)
        else:
            temp_dir = make_request_temp_dir()
            storage_dir = temp_dir

        logger.info('Report request: query=%r, sources=%s, max_pages=%s, file=%s, max_results=%s, since=%s, archive=%s, storage=%s',
                    raw_query, sources, max_pages, file_type, max_results, since, archive, storage_dir)

        # Fetch from providers using query variants (typo tolerance)
        variants = query_variants(raw_query)
        per_provider = max(5, min(100, math.ceil(max_results / max(1, len(sources)))))
        all_records: List[Dict[str, Any]] = []
        for vq in variants:
            for src in sources:
                try:
                    recs = FETCHERS[src](vq, per_provider, since)
                    all_records.extend(recs)
                except Exception as e:
                    logger.warning("Fetcher %s failed for variant %r: %s", src, vq, e)

        if not all_records:
            return jsonify({"error": "No results found from authorized sources with abstracts. Try a different query."}), 404

        # Dedupe, prefer OA, rerank (downloads some PDFs into storage_dir for snippet extraction)
        records = dedupe_records(all_records)
        records = sorted(records, key=lambda r: 0 if r.get("pdf_url") else 1)  # OA first
        ranked = rerank_records(raw_query, records, dest_dir=storage_dir, use_pdf_snippets=True, snippet_limit=8)

        # For image extraction, we may download one or two more PDFs to storage_dir
        one_image = extract_one_image_from_records(ranked, dest_dir=storage_dir)

        # Persist manifest only when archiving (storage_dir is ./local/slug)
        if archive:
            # Ensure PDFs for all ranked with pdf_url are saved into archive folder
            for r in ranked:
                if r.get("pdf_url"):
                    ensure_local_pdf(r, storage_dir)
            write_manifest(storage_dir, raw_query, ranked)

        # Build output in-memory (no file saved on disk for the final PDF/DOCX)
        if file_type == "pdf":
            out_bytes = build_pdf_list_exact_pages(raw_query, ranked, one_image, target_pages=max_pages)
            fname = f'report_{re.sub(r"W+", "_", raw_query)}.pdf'
            resp = send_file(io.BytesIO(out_bytes), mimetype="application/pdf", as_attachment=True, download_name=fname)
        elif file_type == "docx":
            out_bytes = build_docx_list(raw_query, ranked, one_image, max_pages=max_pages)
            fname = f"report_{re.sub(r'W+', '_', raw_query)}.docx"
            resp = send_file(io.BytesIO(out_bytes), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name=fname)
        else:
            return jsonify({"error": "Unsupported file type. Use file=pdf or file=docx."}), 400

        return resp
    finally:
        # If we used a per-request temp dir (archive=0), clean it up now.
        # At this point the response bytes have been created in-memory and sent back.
        if not archive and os.path.isdir(storage_dir):
            cleanup_dir(storage_dir)
        # Delete local subfolder specific to query search if archiving is not enabled
        if not archive and os.path.isdir(os.path.join(LOCAL_ROOT_DIR, safe_slug(raw_query))):
            cleanup_dir(os.path.join(LOCAL_ROOT_DIR, safe_slug(raw_query)))

# ------------------------------------------------------------------------------
# Main
# ------------------------------------------------------------------------------

if __name__ == "__main__":
    desired_debug = os.environ.get("ALLOW_DEBUG", "").lower() in ("1", "true", "yes")
    opts = flask_run_options(desired_debug)
    port = int(os.environ.get("PORT", "8002"))
    logger.info("Starting server on 0.0.0.0:%s (debug=%s)", port, opts.get("debug"))
    app.run(host="0.0.0.0", port=port, **opts)